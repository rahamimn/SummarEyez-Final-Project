try:
    from PIL import Image
except ImportError:
    import Image
import pytesseract
from docx import Document
import pandas as pd 
import numpy as np
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from nltk.tokenize import sent_tokenize,word_tokenize

#will get by arg
file_name = 'test1.jpg'

def extract_hocr_from_jpg():
    hocrFile = open('./output/hocr.html', 'wb')
    hocr = pytesseract.image_to_pdf_or_hocr(file_name, extension='hocr')
    hocrFile.write(hocr)
    return hocr

def extract_text_from_jpg():
    textFile = open('./output/text.text', 'w')
    text = pytesseract.image_to_string(file_name)
    textFile.write(text)
    return text

def extract_tsv_from_jpg():
    tsvFile = open('./output/data.tsv', 'w')
    tsvData = pytesseract.image_to_data(Image.open(file_name))
    tsvFile.write(tsvData)
    return tsvData

def extract_sentences(text):
    sentencesFile = open('./output/sentences.txt', 'w')
    sentences = sent_tokenize(text)
    sentencesFile.writelines(sentences)
    return sentences


def find_index_of_nearest_xy(y_array, x_array, y_point, x_point):
    distance = (y_array - y_point) ** 2 + (x_array - x_point) ** 2
    idx = np.where(distance == distance.min())
    return int(idx[0])

def releate_words_to_senteces(sentences, word_table):
    table_index_loop = 0
    sentencesWeights = [{'text': sent, 'weight': 0} for sent in sentences]
    words_of_sent = [sent.split() for sent in sentences]
    for sent_i, words_sentence in enumerate(words_of_sent):
        for word in words_sentence:
            found = False
            for i in range(table_index_loop,len(word_table)):
                if word_table.loc[i, 'text'] == word:
                #    print('found',i,word, sent_i)
                   word_table.at[i,'sent_i'] = int(sent_i)
                   table_index_loop = i + 1
                   found = True
                   sentencesWeights[sent_i]['weight'] += word_table.at[i,'time_spent']
                   break
            if not found:
                print('notFound',word, sent_i)
    return sentencesWeights

def create_word_dataframe_from_tsv(senteces):
    df = pd.read_table('./output/data.tsv')
    inputXY_df = pd.read_table('./input.tsv')

    df['leftCenter'] = df['left'] + df['width']/2
    df['topCenter'] = df['top'] + df['height']/2
    df['time_spent'] = 0
    df['sent_i'] = -1
    x_array = np.array(df['leftCenter'])
    y_array = np.array(df['topCenter'])

    for index,row in inputXY_df.iterrows():
        x_point = np.array([row['x']])
        y_point = np.array([row['y']])
        idx = find_index_of_nearest_xy(y_array, x_array, y_point, x_point)
        df.at[idx,'time_spent'] = row['time']

    sentencesWeights = releate_words_to_senteces(senteces,df)

    df.to_csv('./output/dataWithFixation.tsv', index=True, sep="\t", na_rep='',
           header=True, index_label=None, mode='w', decimal='.')
    return df, sentencesWeights

def create_docx(word_table_df, sentences_table, mark = None, output_name = 'doc'):
    document = Document()
    paragraph_index = -1
    for index,row in word_table_df.iterrows():
        if(row['par_num'] != paragraph_index):
            p = document.add_paragraph()
            paragraph_index = row['par_num'] 
        # p.font.highlight_color = WD_COLOR_INDEX.YELLOW
        text = row['text']
        if(isinstance(text, str) == False):
            continue
        
        run = p.add_run(text+' ')
        if( mark == 'words' and row['time_spent'] > 0):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if( mark == 'sentences' and sentences_table[row['sent_i']]['weight'] > 0):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.save('./output/'+output_name+'.docx')

def main():
    extract_hocr_from_jpg()
    extract_tsv_from_jpg()
    text = extract_text_from_jpg()
    sentences = extract_sentences(text)
    word_table_df, sentences_table = create_word_dataframe_from_tsv(sentences)
    create_docx(word_table_df,sentences_table, mark='sentences', output_name='docSentences')
    create_docx(word_table_df,sentences_table, mark='words', output_name='docWords')
main()




