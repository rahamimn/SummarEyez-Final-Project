import argparse
import os
from docx import Document
import pandas as pd 
import numpy as np
import nltk
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from nltk.tokenize import sent_tokenize
try:
    from PIL import Image
except ImportError:
    import Image

import pytesseract
import Alg1 
import Alg2 

nltk.data.path.append("/root/nltk_data")

def extract_hocr_from_jpg(file_src):
    hocrFile = open('./output/hocr.html', 'wb')
    hocr = pytesseract.image_to_pdf_or_hocr(file_src, extension='hocr')
    hocrFile.write(hocr)
    return hocr

def extract_text_from_jpg(file_src):
    textFile = open('./output/text.txt', 'w')
    text = pytesseract.image_to_string(file_src)
    textFile.write(text)
    return text

def extract_tsv_from_jpg(file_src):
    tsvFile = open('./output/data.tsv', 'w')
    tsvData = pytesseract.image_to_data(Image.open(file_src), output_type='data.frame')
    tsvFileData = pytesseract.image_to_data(Image.open(file_src))
    tsvFile.write(tsvFileData)
    return tsvData

def extract_sentences(text):
    sentencesFile = open('./output/sentences.txt', 'w')
    sentences = sent_tokenize(text)
    sentences = [sent.replace('\n',' ') for sent in sentences]
    sentencesFile.writelines(sentences)
    return sentences


def find_index_of_nearest_xy(y_array, x_array, y_point, x_point):
    distance = (y_array - y_point) ** 2 + (x_array - x_point) ** 2
    idx = np.where(distance == distance.min())
    return int(idx[0])

def create_sentences_table(word_ocr, sentences):
    table_index_loop = 0
    par_num = 0
    sentencesWeights = [{'text': sent, 'par_num': -1, 'word_count': 0, 'char_count': len(sent) } for sent in sentences]
    words_of_sent = [sent.split() for sent in sentences]
    for sent_num, words_sentence in enumerate(words_of_sent):
        #hack to find if the right paragraph
        first_word = True 
        empty_words_counter = 0
        sentencesWeights[sent_num]['word_count'] += len(words_sentence)
        for word in words_sentence:
            found = False
            for i in range(table_index_loop,len(word_ocr)):
                if pd.isna(word_ocr.loc[i, 'text']):
                    empty_words_counter += 1
                    
                if word_ocr.loc[i, 'text'] == word:
                #    print('found',i,word, sent_num)
                    if empty_words_counter > 1 and first_word :
                        par_num += 1
                    first_word = False
                    word_ocr.at[i,'sent_num'] = int(sent_num)
                    word_ocr.at[i,'par_num'] = par_num
                    table_index_loop = i + 1
                    if sentencesWeights[sent_num]['par_num'] == -1:
                        sentencesWeights[sent_num]['par_num'] = par_num
                    found = True
                    break
            if not found:
                print('notFound',word, sent_num)
    #remove empty text
    word_ocr = word_ocr[pd.isna(word_ocr.text) == False ]
    word_ocr = word_ocr.reset_index()
    return word_ocr, pd.DataFrame(sentencesWeights)


def pre_process(word_ocr, sentences):
    #delete non important colomns
    #important {text,par_num}
    word_ocr['center_x'] = word_ocr['left'] + word_ocr['width'] / 2 
    word_ocr['center_y'] = word_ocr['top'] + word_ocr['height'] / 2
    word_ocr['sent_num'] = -1

    word_ocr = word_ocr[['par_num','text','center_x','center_y','sent_num']]

    #create sent table 
    word_ocr, sentences_table = create_sentences_table(word_ocr, sentences)
    

    sentences_table.to_csv('./output/sent_table_base.tsv', index=True, sep="\t", na_rep='',
        header=True, index_label=None, mode='w', decimal='.')
    word_ocr.to_csv('./output/word_ocr.tsv', index=True, sep="\t", na_rep='', 
        header=True, index_label=None, mode='w', decimal='.')
    return word_ocr, sentences_table

def create_docx_words(word_table, output_name = 'doc'):
    document = Document()
    paragraph_index = -1
    for index,row in word_table.iterrows():
        if(row['par_num'] != paragraph_index):
            p = document.add_paragraph()
            paragraph_index = row['par_num']
        text = row['text']
        run = p.add_run(text+' ')
        if(row['weight'] > 0):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.save('./docs/'+output_name+'_word.docx')

def create_docx_sents(sentences_table, output_name = 'doc', threshold = 0):
    document = Document()
    paragraph_index = -1
    for index,row in sentences_table.iterrows():
        if(row['par_num'] != paragraph_index):
            p = document.add_paragraph()
            paragraph_index = row['par_num'] 
        # p.font.highlight_color = WD_COLOR_INDEX.YELLOW
        text = row['text']
        run = p.add_run(text+' ')
        if(row['weight'] > threshold):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    document.save('./docs/'+output_name+'_sent.docx')

def calculate_weight(fixations, word_table, sentences_table, filename):
    x_array = np.array(word_table['center_x'])
    y_array = np.array(word_table['center_y'])

    word_table['weight'] = 0
    sentences_table['weight'] = 0

    for index,row in fixations.iterrows():
        x_point = np.array([row['Fixation Position X [px]']]) #real name
        y_point = np.array([row['Fixation Position Y [px]']]) #real name
        idx = find_index_of_nearest_xy(y_array, x_array, y_point, x_point) 
        word_table.at[idx,'weight'] += int(row['Event Duration [ms]']) #real name
        sent_num =  word_table.loc[idx, 'sent_num']
        if(sent_num > -1):
            sentences_table.at[sent_num,'weight']  +=  int(row['Event Duration [ms]'])

    word_table = word_table[['text', 'sent_num', 'par_num', 'weight']]
    return word_table, sentences_table


def handleAutomaticAlg(text):
    alg1_sent_table = pd.DataFrame(Alg1.run_summarization(text))
    alg1_sent_table['par_num'] = 1

    sortedAlg1 = alg1_sent_table.sort_values(['weight'], ascending=True)
    val = 1

    for index,row in sortedAlg1.iterrows():
        alg1_sent_table.loc[index,'weight'] = val
        val+=1

    alg2_sent_table = pd.DataFrame(Alg2.calculate_weights(text))
    alg2_sent_table['par_num'] = 1

    alg3_sent_table = pd.read_csv('Alg3.csv')
    alg3_sent_table['par_num'] = 1

    for sent_table in [alg1_sent_table, alg2_sent_table, alg3_sent_table] :
        sent_table.sort_values(by=['weight']) 
        max_value = sent_table['weight'].max()
        min_value = sent_table['weight'].min()
        sent_table['weight'] = (sent_table['weight'] - min_value) / (max_value - min_value)
    create_docx_sents(alg1_sent_table, output_name='Alg1-Sents', threshold = 0.8)
    create_docx_sents(alg2_sent_table, output_name='Alg2-Sents', threshold = 0.8)
    create_docx_sents(alg3_sent_table, output_name='Alg3-Sents', threshold = 0.8)

def handleFixations(fixations, word_table, sentences_table, filename):
    word_table, sent_table = calculate_weight(fixations, word_table, sentences_table, filename)

    max_value = sent_table['weight'].max()
    min_value = sent_table['weight'].min()
    sent_table['weight'] = (sent_table['weight'] - min_value) / (max_value - min_value)

    word_table.to_csv('./output/'+filename+'-word_table.tsv', index=True, sep="\t", na_rep='',
        header=True, index_label=None, mode='w', decimal='.')

    sent_table.to_csv('./output/'+filename+'-sent_table.tsv', index=True, sep="\t", na_rep='',
        header=True, index_label=None, mode='w', decimal='.')

    create_docx_words(word_table, output_name=filename+'-Words')
    create_docx_sents(sentences_table, output_name=filename+'-Sents')

def main():
    extract_hocr_from_jpg('test2.jpg')
    word_ocr = extract_tsv_from_jpg('test2.jpg')
    text = extract_text_from_jpg('test2.jpg')
    sentences = extract_sentences(text)

    word_ocr, sentences_table_base = pre_process(word_ocr, sentences)
    handleAutomaticAlg(text)
    #for each File

    for filename in os.listdir('./tests'):
        fixations = pd.read_csv('./tests/'+filename)
        handleFixations(fixations, word_ocr.copy(), sentences_table_base.copy(), filename)

if __name__ == "__main__":
    main()



