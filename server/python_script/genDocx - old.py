from docx import Document
from docx.enum.text import WD_COLOR_INDEX

#get tesseract_words table and seneteces/words grade (maybe plural for layering)
#generate docx file 


def create_docx_words(word_table, output_name = 'doc'):
    document = Document()
    paragraph_index = -1
    for index,row in word_table.iterrows():
        if(row['par_num'] != paragraph_index):
            p = document.add_paragraph()
            paragraph_index = row['par_num']
        text = row['text']
        run = p.add_run(text+' ')
        if(row['weight'] > 0):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.save('./docs/'+output_name+'_word.docx')

def create_docx_sents(sentences_table, output_name = 'doc', threshold = 0):
    document = Document()
    paragraph_index = -1
    for index,row in sentences_table.iterrows():
        if(row['par_num'] != paragraph_index):
            p = document.add_paragraph()
            paragraph_index = row['par_num'] 
        # p.font.highlight_color = WD_COLOR_INDEX.YELLOW
        text = row['text']
        run = p.add_run(text+' ')
        if(row['weight'] > threshold):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    document.save('./docs/'+output_name+'_sent.docx')

